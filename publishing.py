import docx

import os
class DocWriter:

    def __init__(self) -> None:
        self.doc = docx.Document()

    def write_doc(self, book, chapter_dict, title):

        self.doc.add_heading(title, 0)

        for chapter, paragraphs_list in book.items():

            description = chapter_dict[chapter]
            chapter_name = '{}: {}'.format(
                chapter.strip(), description.strip()
            )

            self.doc.add_heading(chapter_name, 1)

            text = '\n\n'.join(paragraphs_list)
            self.doc.add_paragraph(text)
        file_path = os.path.join(os.path.dirname(__file__),'docs', 'book.docx')
        self.doc.save(file_path)
